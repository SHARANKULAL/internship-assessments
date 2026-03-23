"""
step2_dqa_assessment.py
Reads the 6 raw CSVs, runs quality checks, and produces:
  - issues_log.xlsx  (detailed issue log)
  - DQA_Issues_Report.docx  (Word documentation)
"""

import pandas as pd
import os, re
from datetime import date
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RAW   = r'C:\Users\shara\OneDrive\Desktop\Aptitude\week3\dataset\raw'
OUT   = r'C:\Users\shara\OneDrive\Desktop\Aptitude\week3\dataset'
os.makedirs(OUT, exist_ok=True)

# ─── Load tables ──────────────────────────────────────────────────────────────
print('📂  Loading raw datasets ...')
learners  = pd.read_csv(f'{RAW}/learners.csv',          dtype=str).fillna('')
courses   = pd.read_csv(f'{RAW}/courses.csv',           dtype=str).fillna('')
asmts     = pd.read_csv(f'{RAW}/assessments.csv',       dtype=str).fillna('')
enrolls   = pd.read_csv(f'{RAW}/course_enrollments.csv',dtype=str).fillna('')
txns      = pd.read_csv(f'{RAW}/transactions.csv',      dtype=str).fillna('')
calendar  = pd.read_csv(f'{RAW}/academic_calendar.csv', dtype=str).fillna('')

tables = {
    'learners':           learners,
    'courses':            courses,
    'assessments':        asmts,
    'course_enrollments': enrolls,
    'transactions':       txns,
    'academic_calendar':  calendar,
}

# ─── Issue collector ──────────────────────────────────────────────────────────
issues = []   # list of dicts

def log(table, row_idx, field, issue_type, details, severity, resolution):
    issues.append({
        'Table':      table,
        'Row #':      row_idx,
        'Field':      field,
        'Issue Type': issue_type,
        'Details':    details,
        'Severity':   severity,
        'Resolution': resolution,
    })

def safe_date(s):
    try:
        return date.fromisoformat(str(s).strip())
    except Exception:
        return None

# ─── LEARNERS Checks ──────────────────────────────────────────────────────────
print('🔍  Checking learners ...')
REQ = ['learner_id','email','enrollment_date']
EMAIL_RE = re.compile(r'^[^@\s]+@[^@\s]+\.[^@\s]+$')
valid_statuses = {'active','inactive','pending'}

for idx, row in learners.iterrows():
    rn = idx + 2
    # Null required fields
    for f in REQ:
        if not str(row.get(f,'')).strip():
            log('learners', rn, f, 'Missing Value', f'Required field "{f}" is empty', 'High',
                f'Populate "{f}" from source system; reject record if unresolvable')
    # Missing first_name
    if not str(row.get('first_name','')).strip():
        log('learners', rn, 'first_name', 'Missing Value','first_name is blank','Medium',
            'Look up learner profile by ID; flag for manual review')
    # Email format
    em = str(row.get('email','')).strip()
    if em and not EMAIL_RE.match(em):
        log('learners', rn, 'email', 'Invalid Format', f'Email "{em}" fails format check', 'High',
            'Correct email or mark learner as uncontactable; re-verify via registration source')
    # Future DOB
    dob = safe_date(row.get('date_of_birth',''))
    if dob and dob > date.today():
        log('learners', rn, 'date_of_birth', 'Invalid Value', f'DOB {dob} is in the future', 'High',
            'Correct to actual birth date; likely data entry transposition error')
    # Inconsistent status
    st = str(row.get('status','')).strip()
    if st and st.lower() in valid_statuses and st not in {'Active','Inactive','Pending'}:
        log('learners', rn, 'status', 'Inconsistent Format',
            f'Status "{st}" inconsistent casing (expected Title Case)', 'Low',
            'Standardise to Title Case: Active / Inactive / Pending')
    # Early enrollment
    ed = safe_date(row.get('enrollment_date',''))
    if ed and ed < date(2021, 1, 1):
        log('learners', rn, 'enrollment_date', 'Invalid Value',
            f'Enrollment date {ed} predates platform launch (Jan 2021)', 'Medium',
            'Verify against registration logs; correct or mark for manual review')

# Duplicate learner_id
dup_lid = learners[learners['learner_id'].duplicated(keep=False)]
for idx in dup_lid.index:
    log('learners', idx+2, 'learner_id', 'Duplicate Key',
        f'learner_id "{dup_lid.loc[idx,"learner_id"]}" is duplicated', 'Critical',
        'Retain earliest record; merge or delete duplicate; reassign foreign keys')

# Duplicate email
dup_em = learners[(learners['email'] != '') & learners['email'].duplicated(keep=False)]
for idx in dup_em.index:
    log('learners', idx+2, 'email', 'Duplicate Value',
        f'Email "{dup_em.loc[idx,"email"]}" appears more than once', 'High',
        'Merge duplicate learner accounts or correct email entry error')

# Full duplicate rows
dup_rows = learners[learners.duplicated(keep=False)]
seen = set()
for idx in dup_rows.index:
    key = tuple(dup_rows.loc[idx])
    if key not in seen:
        seen.add(key)
    else:
        log('learners', idx+2, 'ALL', 'Duplicate Row',
            'Row is an exact duplicate of a previous record', 'Critical',
            'Delete duplicate record; preserve the earliest created row')

# Phone with non-numeric
for idx, row in learners.iterrows():
    ph = str(row.get('phone','')).strip()
    if ph and re.search(r'[A-Za-z]', ph):
        log('learners', idx+2, 'phone', 'Invalid Format',
            f'Phone "{ph}" contains alphabetic characters', 'Medium',
            'Clean to numeric digits only and verify against source')

# ─── COURSES Checks ──────────────────────────────────────────────────────────
print('🔍  Checking courses ...')
for idx, row in courses.iterrows():
    rn = idx + 2
    if not str(row.get('course_title','')).strip():
        log('courses', rn, 'course_title', 'Missing Value', 'course_title is blank', 'High',
            'Retrieve title from course catalogue; do not publish without title')
    if not str(row.get('instructor_id','')).strip():
        log('courses', rn, 'instructor_id', 'Missing Value', 'instructor_id is blank', 'Medium',
            'Assign an instructor before course activation')
    try:
        dur = float(str(row.get('duration_hours',0)))
        if dur <= 0:
            log('courses', rn, 'duration_hours', 'Invalid Value',
                f'duration_hours = {dur} is ≤ 0', 'High',
                'Set to correct duration; zero/negative duration is logically invalid')
    except: pass
    try:
        pr = float(str(row.get('price',0)))
        if pr < 0:
            log('courses', rn, 'price', 'Invalid Value',
                f'price = {pr} is negative', 'High',
                'Correct to actual price; negative price likely data entry error')
    except: pass
    sd = safe_date(row.get('start_date',''))
    ed = safe_date(row.get('end_date',''))
    if sd and ed and ed < sd:
        log('courses', rn, 'end_date', 'Date Logic Error',
            f'end_date {ed} is before start_date {sd}', 'Critical',
            'Swap dates or correct to actual course schedule')
    st = str(row.get('status','')).strip().lower()
    valid_course_st = {'published','draft','archived'}
    if st and st not in valid_course_st:
        log('courses', rn, 'status', 'Invalid Categorical',
            f'Status "{row["status"]}" not in valid values', 'Low',
            'Correct to Published / Draft / Archived')

# ─── ASSESSMENTS Checks ──────────────────────────────────────────────────────
print('🔍  Checking assessments ...')
valid_cids = set(courses['course_id'].tolist())
for idx, row in asmts.iterrows():
    rn = idx + 2
    if not str(row.get('assessment_type','')).strip():
        log('assessments', rn, 'assessment_type', 'Missing Value',
            'assessment_type is blank', 'Medium',
            'Classify as Quiz / Assignment / Project / Exam')
    try:
        mx = float(str(row.get('max_score',0)))
        ps = float(str(row.get('passing_score',0)))
        if mx <= 0:
            log('assessments', rn, 'max_score', 'Invalid Value',
                f'max_score = {mx} must be > 0', 'High',
                'Set to correct maximum marks value')
        if ps > mx:
            log('assessments', rn, 'passing_score', 'Business Rule Violation',
                f'passing_score {ps} > max_score {mx}', 'Critical',
                'passing_score must be ≤ max_score; correct assessment configuration')
    except: pass
    dd = safe_date(row.get('due_date',''))
    cd = safe_date(row.get('created_date',''))
    if dd and cd and dd < cd:
        log('assessments', rn, 'due_date', 'Date Logic Error',
            f'due_date {dd} is before created_date {cd}', 'High',
            'Correct due_date to a date after assessment creation')
    cid = str(row.get('course_id','')).strip()
    if cid and cid not in valid_cids:
        log('assessments', rn, 'course_id', 'Referential Integrity',
            f'course_id "{cid}" does not exist in courses table', 'Critical',
            'Link to valid course_id or delete orphan assessment')

dup_asmt = asmts[asmts['assessment_id'].duplicated(keep=False)]
for idx in dup_asmt.index:
    log('assessments', idx+2, 'assessment_id', 'Duplicate Key',
        f'assessment_id "{dup_asmt.loc[idx,"assessment_id"]}" is duplicated', 'Critical',
        'Retain active record; delete or renumber duplicate')

# ─── ENROLLMENTS Checks ──────────────────────────────────────────────────────
print('🔍  Checking course_enrollments ...')
valid_lids = set(learners['learner_id'].drop_duplicates().tolist())
for idx, row in enrolls.iterrows():
    rn = idx + 2
    if not str(row.get('enrollment_date','')).strip():
        log('course_enrollments', rn, 'enrollment_date', 'Missing Value',
            'enrollment_date is blank', 'High',
            'Derive from system logs or mark as unknown; required for analytics')
    lid = str(row.get('learner_id','')).strip()
    cid = str(row.get('course_id','')).strip()
    if lid and lid not in valid_lids:
        log('course_enrollments', rn, 'learner_id', 'Referential Integrity',
            f'learner_id "{lid}" not found in learners table', 'Critical',
            'Validate against learners; delete orphan or correct typo')
    if cid and cid not in valid_cids:
        log('course_enrollments', rn, 'course_id', 'Referential Integrity',
            f'course_id "{cid}" not found in courses table', 'Critical',
            'Validate against courses; delete orphan or correct typo')
    ed2 = safe_date(row.get('enrollment_date',''))
    cd2 = safe_date(row.get('completion_date',''))
    if ed2 and cd2 and cd2 < ed2:
        log('course_enrollments', rn, 'completion_date', 'Date Logic Error',
            f'completion_date {cd2} is before enrollment_date {ed2}', 'High',
            'Correct date sequence; swap if data entry error')
    try:
        gr = float(str(row.get('grade',0)))
        if gr < 0 or gr > 100:
            log('course_enrollments', rn, 'grade', 'Out of Range',
                f'grade = {gr} outside valid 0–100 range', 'High',
                'Cap at 100 or investigate; grades above 100 suggest data corruption')
    except: pass

dup_enr = enrolls[enrolls['enrollment_id'].duplicated(keep=False)]
for idx in dup_enr.index:
    log('course_enrollments', idx+2, 'enrollment_id', 'Duplicate Key',
        f'enrollment_id "{dup_enr.loc[idx,"enrollment_id"]}" duplicated', 'Critical',
        'Delete duplicate; keep earliest timestamp record')

dup_enr_row = enrolls[enrolls.duplicated(keep=False)]
seen2 = set()
for idx in dup_enr_row.index:
    key = tuple(enrolls.loc[idx])
    if key not in seen2:
        seen2.add(key)
    else:
        log('course_enrollments', idx+2, 'ALL', 'Duplicate Row',
            'Entire enrollment row duplicated', 'Critical',
            'Remove duplicate enrollment row')

# ─── TRANSACTIONS Checks ─────────────────────────────────────────────────────
print('🔍  Checking transactions ...')
valid_currencies = {'INR','USD','EUR','GBP','AED'}
for idx, row in txns.iterrows():
    rn = idx + 2
    if not str(row.get('payment_method','')).strip():
        log('transactions', rn, 'payment_method', 'Missing Value',
            'payment_method is blank', 'Medium',
            'Derive from payment gateway logs; mark as Unknown if unavailable')
    if not str(row.get('transaction_date','')).strip():
        log('transactions', rn, 'transaction_date', 'Missing Value',
            'transaction_date is blank', 'High',
            'Retrieve from payment processor timestamp; required for financial reporting')
    try:
        amt = float(str(row.get('amount',0)))
        if amt < 0:
            log('transactions', rn, 'amount', 'Invalid Value',
                f'amount = {amt} is negative', 'Critical',
                'If refund, create separate refund record; correct sign')
        if amt == 0:
            log('transactions', rn, 'amount', 'Invalid Value',
                f'amount = 0; zero-value transaction requires justification', 'Medium',
                'Verify if scholarship/waiver; link to waiver record or delete')
        if amt > 99999:
            log('transactions', rn, 'amount', 'Outlier / Suspicious Value',
                f'amount = {amt} unusually large (> ₹99,999)', 'Medium',
                'Verify against payment gateway receipt; possible data entry error')
    except: pass
    lid2 = str(row.get('learner_id','')).strip()
    if lid2 and lid2 not in valid_lids:
        log('transactions', rn, 'learner_id', 'Referential Integrity',
            f'learner_id "{lid2}" not in learners table', 'Critical',
            'Correct learner_id or investigate orphan transaction')
    cur = str(row.get('currency','')).strip().upper()
    if cur and cur not in valid_currencies:
        log('transactions', rn, 'currency', 'Invalid Categorical',
            f'currency "{row["currency"]}" not a recognised ISO code', 'Medium',
            'Correct to ISO 4217 code (INR / USD / EUR etc.)')
    st2 = str(row.get('status','')).strip()
    valid_txn_st = {'Completed','Failed','Refunded','Pending'}
    if st2 and st2 not in valid_txn_st:
        log('transactions', rn, 'status', 'Invalid Categorical',
            f'status "{st2}" not in valid values', 'Medium',
            'Correct to Completed / Failed / Refunded / Pending')

dup_txn = txns[txns['transaction_id'].duplicated(keep=False)]
for idx in dup_txn.index:
    log('transactions', idx+2, 'transaction_id', 'Duplicate Key',
        f'transaction_id "{dup_txn.loc[idx,"transaction_id"]}" duplicated', 'Critical',
        'Keep original; mark duplicate as voided or delete')

dup_txn_row = txns[txns.duplicated(keep=False)]
seen3 = set()
for idx in dup_txn_row.index:
    key = tuple(txns.loc[idx])
    if key not in seen3:
        seen3.add(key)
    else:
        log('transactions', idx+2, 'ALL', 'Duplicate Row',
            'Entire transaction row duplicated', 'Critical',
            'Delete duplicate transaction record')

# ─── CALENDAR Checks ─────────────────────────────────────────────────────────
print('🔍  Checking academic_calendar ...')
for idx, row in calendar.iterrows():
    rn = idx + 2
    if not str(row.get('term_name','')).strip():
        log('academic_calendar', rn, 'term_name', 'Missing Value',
            'term_name is blank', 'Medium',
            'Assign academic term name from institutional calendar')
    if not str(row.get('start_date','')).strip():
        log('academic_calendar', rn, 'start_date', 'Missing Value',
            'start_date is blank', 'High',
            'Set from official academic schedule')
    sd2 = safe_date(row.get('start_date',''))
    ed2 = safe_date(row.get('end_date',''))
    if sd2 and ed2 and ed2 < sd2:
        log('academic_calendar', rn, 'end_date', 'Date Logic Error',
            f'end_date {ed2} is before start_date {sd2}', 'Critical',
            'Correct date sequence; likely year was swapped in data entry')

dup_cal = calendar[calendar['calendar_id'].duplicated(keep=False)]
for idx in dup_cal.index:
    log('academic_calendar', idx+2, 'calendar_id', 'Duplicate Key',
        f'calendar_id "{dup_cal.loc[idx,"calendar_id"]}" is duplicated', 'Critical',
        'Renumber duplicate calendar ID or merge overlapping terms')

# Overlapping term check
cal_clean = calendar[(calendar['start_date'] != '') & (calendar['end_date'] != '')].copy()
cal_clean['_sd'] = cal_clean['start_date'].apply(safe_date)
cal_clean['_ed'] = cal_clean['end_date'].apply(safe_date)
cal_clean = cal_clean.dropna(subset=['_sd','_ed'])
cal_clean_list = cal_clean.reset_index(drop=True)
for i in range(len(cal_clean_list)):
    for j in range(i+1, len(cal_clean_list)):
        a_sd = cal_clean_list.loc[i, '_sd']
        a_ed = cal_clean_list.loc[i, '_ed']
        b_sd = cal_clean_list.loc[j, '_sd']
        b_ed = cal_clean_list.loc[j, '_ed']
        a_name = cal_clean_list.loc[i, 'term_name']
        b_name = cal_clean_list.loc[j, 'term_name']
        if a_sd and b_sd and a_ed and b_ed:
            if not (a_ed < b_sd or b_ed < a_sd):
                log('academic_calendar', j+2, 'start_date/end_date',
                    'Overlapping Terms',
                    f'Term "{b_name}" overlaps with "{a_name}"', 'High',
                    'Adjust term boundaries so no two active terms overlap')

# ─── Summary ─────────────────────────────────────────────────────────────────
print(f'\n📊  Total issues found: {len(issues)}')
sev = {}
for iss in issues:
    sev[iss['Severity']] = sev.get(iss['Severity'], 0) + 1
for k, v in sorted(sev.items()):
    print(f'   {k}: {v}')

# ═══════════════════════════════════════════════════════════════════════════════
# WRITE EXCEL ISSUES LOG
# ═══════════════════════════════════════════════════════════════════════════════
print('\n📊  Writing issues_log.xlsx ...')
df = pd.DataFrame(issues)
sev_order = {'Critical':0,'High':1,'Medium':2,'Low':3}
df['_sev_order'] = df['Severity'].map(sev_order)
df.sort_values(['_sev_order','Table'], inplace=True)
df.drop(columns=['_sev_order'], inplace=True)

wb = Workbook()
ws = wb.active
ws.title = 'Issues Log'

# Header style
HDR_FILL = PatternFill('solid', fgColor='1A509C')
HDR_FONT = Font(name='Calibri', bold=True, color='FFFFFF', size=11)
BODY_FONT = Font(name='Calibri', size=10)
ALT_FILL = PatternFill('solid', fgColor='E9F1FB')
thin = Side(style='thin', color='CCDDEE')
BORDER = Border(left=thin, right=thin, top=thin, bottom=thin)
SEV_COLORS = {
    'Critical': 'FF4444',
    'High':     'FF8C00',
    'Medium':   'F5A623',
    'Low':      '00878A',
}

columns = list(df.columns)
col_widths = [22,7,22,26,55,12,55]

for ci, col in enumerate(columns, 1):
    c = ws.cell(row=1, column=ci, value=col)
    c.fill = HDR_FILL
    c.font = HDR_FONT
    c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    c.border = BORDER
    if ci <= len(col_widths):
        ws.column_dimensions[get_column_letter(ci)].width = col_widths[ci-1]
ws.row_dimensions[1].height = 28

for ri, (_, row) in enumerate(df.iterrows(), 2):
    is_alt = ri % 2 == 0
    for ci, col in enumerate(columns, 1):
        val = row[col]
        c = ws.cell(row=ri, column=ci, value=val)
        c.font = BODY_FONT
        c.alignment = Alignment(vertical='top', wrap_text=True)
        c.border = BORDER
        if col == 'Severity':
            color = SEV_COLORS.get(str(val), 'CCCCCC')
            c.fill = PatternFill('solid', fgColor=color)
            c.font = Font(name='Calibri', size=10, bold=True, color='FFFFFF')
            c.alignment = Alignment(horizontal='center', vertical='top')
        elif is_alt and col != 'Severity':
            c.fill = ALT_FILL

# Summary sheet
ws2 = wb.create_sheet('Summary')
ws2.column_dimensions['A'].width = 28
ws2.column_dimensions['B'].width = 15
ws2.column_dimensions['C'].width = 18

sum_headers = ['Category','Count','% of Total']
for ci, h in enumerate(sum_headers, 1):
    c = ws2.cell(row=1, column=ci, value=h)
    c.fill = HDR_FILL; c.font = HDR_FONT
    c.alignment = Alignment(horizontal='center'); c.border = BORDER

# By severity
row_idx = 2
ws2.cell(row=row_idx, column=1, value='── BY SEVERITY ──').font = Font(bold=True, size=10, color='1A509C')
row_idx += 1
for sev_name in ['Critical','High','Medium','Low']:
    cnt = len(df[df['Severity']==sev_name])
    pct = f'{cnt/len(df)*100:.1f}%'
    for ci, v in enumerate([sev_name, cnt, pct], 1):
        c = ws2.cell(row=row_idx, column=ci, value=v)
        c.font = BODY_FONT; c.border = BORDER
        c.fill = PatternFill('solid', fgColor='E9F1FB') if row_idx%2==0 else PatternFill('solid', fgColor='FFFFFF')
    row_idx += 1

# By table
row_idx += 1
ws2.cell(row=row_idx, column=1, value='── BY TABLE ──').font = Font(bold=True, size=10, color='1A509C')
row_idx += 1
for tbl in df['Table'].unique():
    cnt = len(df[df['Table']==tbl])
    pct = f'{cnt/len(df)*100:.1f}%'
    for ci, v in enumerate([tbl, cnt, pct], 1):
        c = ws2.cell(row=row_idx, column=ci, value=v)
        c.font = BODY_FONT; c.border = BORDER
        c.fill = PatternFill('solid', fgColor='E9F1FB') if row_idx%2==0 else PatternFill('solid', fgColor='FFFFFF')
    row_idx += 1

# By issue type
row_idx += 1
ws2.cell(row=row_idx, column=1, value='── BY ISSUE TYPE ──').font = Font(bold=True, size=10, color='1A509C')
row_idx += 1
for itype, grp in df.groupby('Issue Type'):
    cnt = len(grp)
    pct = f'{cnt/len(df)*100:.1f}%'
    for ci, v in enumerate([itype, cnt, pct], 1):
        c = ws2.cell(row=row_idx, column=ci, value=v)
        c.font = BODY_FONT; c.border = BORDER
        c.fill = PatternFill('solid', fgColor='E9F1FB') if row_idx%2==0 else PatternFill('solid', fgColor='FFFFFF')
    row_idx += 1

wb.save(f'{OUT}/issues_log.xlsx')
print(f'  ✅  issues_log.xlsx saved')

# ═══════════════════════════════════════════════════════════════════════════════
# WRITE WORD ISSUES REPORT
# ═══════════════════════════════════════════════════════════════════════════════
print('📄  Writing DQA_Issues_Report.docx ...')

NAVY  = RGBColor(0x1A, 0x50, 0x9C)
TEAL  = RGBColor(0x00, 0x87, 0x8A)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GRAY  = RGBColor(0x55, 0x65, 0x7A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def shade_cell2(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def doc_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(18); r.font.bold = True
    r.font.color.rgb = NAVY
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0); rule.paragraph_format.space_after = Pt(8)
    rr = rule.add_run('─' * 85); rr.font.size = Pt(5); rr.font.color.rgb = TEAL

def doc_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = TEAL

def doc_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = DARK

def doc_bullet(text, sym='▸'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f'{sym}  {text}')
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = DARK

def issues_table(df_sub):
    cols = ['Row #','Field','Issue Type','Details','Severity','Resolution']
    col_w = [0.6, 1.2, 1.5, 2.4, 0.85, 2.8]
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    hr = tbl.rows[0]
    for ci, hdr in enumerate(cols):
        c = hr.cells[ci]
        shade_cell2(c, '1A509C')
        c.width = Inches(col_w[ci])
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(hdr)
        r.font.bold = True; r.font.color.rgb = WHITE
        r.font.name = 'Calibri'; r.font.size = Pt(9)
    SEV_HEX = {'Critical':'FF4444','High':'FF8C00','Medium':'F5A623','Low':'00878A'}
    for ri, (_, row) in enumerate(df_sub.iterrows()):
        dr = tbl.add_row()
        bg = 'E9F1FB' if ri%2==0 else 'FFFFFF'
        for ci, col in enumerate(cols):
            c = dr.cells[ci]
            val = str(row.get(col,''))
            if col == 'Severity':
                shade_cell2(c, SEV_HEX.get(val,'CCCCCC'))
                c.width = Inches(col_w[ci])
                p = c.paragraphs[0]
                r = p.add_run(val)
                r.font.bold = True; r.font.color.rgb = WHITE
                r.font.name = 'Calibri'; r.font.size = Pt(9)
            else:
                shade_cell2(c, bg)
                c.width = Inches(col_w[ci])
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
                r = p.add_run(val)
                r.font.name = 'Calibri'; r.font.size = Pt(9); r.font.color.rgb = DARK
    doc.add_paragraph()

# ── Cover ─────────────────────────────────────────────────────────────────────
cp = doc.add_paragraph()
cp.paragraph_format.space_before = Pt(48); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cp.add_run('DATA QUALITY ASSESSMENT\nISSUES DOCUMENTATION')
cr.font.name = 'Calibri'; cr.font.size = Pt(26); cr.font.bold = True; cr.font.color.rgb = NAVY
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run('EdTech Platform — Raw Dataset Quality Report')
sr.font.name = 'Calibri'; sr.font.size = Pt(14); sr.font.italic = True; sr.font.color.rgb = TEAL
rp = doc.add_paragraph(); rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rp.paragraph_format.space_before = Pt(6)
rr2 = rp.add_run('━' * 42); rr2.font.color.rgb = RGBColor(0xF5,0xA6,0x23)
mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = mp.add_run(f'Total Issues: {len(df)}  |  Critical: {sev.get("Critical",0)}  |  High: {sev.get("High",0)}  |  March 2026')
mr.font.name = 'Calibri'; mr.font.size = Pt(11); mr.font.color.rgb = GRAY
doc.add_page_break()

# ── 1. Executive Summary ──────────────────────────────────────────────────────
doc_h1('1. Executive Summary')
doc_body(
    f'A comprehensive Data Quality Assessment was conducted on the EdTech platform raw dataset '
    f'comprising 6 tables and approximately 225 records. The automated quality assessment engine '
    f'identified a total of {len(df)} distinct data quality issues spanning all six tables.'
)
doc_body('Issue severity breakdown:')
for sev_name in ['Critical','High','Medium','Low']:
    cnt = len(df[df['Severity']==sev_name])
    doc_bullet(f'{sev_name}: {cnt} issues ({cnt/len(df)*100:.1f}% of total)')
doc_body('Issue type categories identified:')
for itype, grp in df.groupby('Issue Type'):
    doc_bullet(f'{itype}: {len(grp)} occurrences')

# ── 2. Dataset Profile ────────────────────────────────────────────────────────
doc_h1('2. Dataset Profile')
doc_body('The following tables were assessed:')
table_info = [
    ('learners',           50,  9, 'Core learner profiles'),
    ('courses',            20,  9, 'Course catalogue'),
    ('assessments',        25,  7, 'Quizzes, assignments, exams'),
    ('course_enrollments', 70,  7, 'Learner–course linkage (Enrolments)'),
    ('transactions',       50,  8, 'Payment & billing records (Transaction Table)'),
    ('academic_calendar',  10,  7, 'Term dates and schedules (Calendar Table)'),
]
cols = ['Table','Records','Fields','Description']
col_w2 = [2.0, 0.9, 0.7, 3.7]
tbl2 = doc.add_table(rows=1, cols=4)
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl2.style = 'Table Grid'
hr2 = tbl2.rows[0]
for ci, h in enumerate(cols):
    c = hr2.cells[ci]
    shade_cell2(c, '1A509C')
    c.width = Inches(col_w2[ci])
    p = c.paragraphs[0]
    r = p.add_run(h); r.font.bold = True; r.font.color.rgb = WHITE
    r.font.name = 'Calibri'; r.font.size = Pt(10)
for ri, (tname, recs, flds, desc) in enumerate(table_info):
    dr2 = tbl2.add_row()
    bg = 'E9F1FB' if ri%2==0 else 'FFFFFF'
    for ci, v in enumerate([tname, recs, flds, desc]):
        c = dr2.cells[ci]; shade_cell2(c, bg)
        c.width = Inches(col_w2[ci])
        p = c.paragraphs[0]; r = p.add_run(str(v))
        r.font.name = 'Calibri'; r.font.size = Pt(10); r.font.color.rgb = DARK
doc.add_paragraph()

# ── 3–8. Per-Table Issue Details ──────────────────────────────────────────────
for i, tname in enumerate(['learners','courses','assessments',
                             'course_enrollments','transactions','academic_calendar'], 3):
    df_tbl = df[df['Table'] == tname]
    doc_h1(f'{i}. Issues — {tname} Table')
    cnt_c = len(df_tbl[df_tbl['Severity']=='Critical'])
    cnt_h = len(df_tbl[df_tbl['Severity']=='High'])
    doc_body(f'Total issues: {len(df_tbl)}  |  Critical: {cnt_c}  |  High: {cnt_h}')
    if len(df_tbl):
        issues_table(df_tbl)
    else:
        doc_body('No issues detected in this table.')

# ── 9. Resolution Priority ────────────────────────────────────────────────────
doc_h1('9. Recommended Resolution Priority')
doc_body('Prioritise issue resolution in the following order:')
prios = [
    ('P0 — Immediate (Pre-Analytics)',  'Critical',
     'Duplicate primary keys, referential integrity failures, date logic errors in core tables'),
    ('P1 — This Sprint (Within 1 Week)','High',
     'Missing required fields, invalid values, negative amounts, invalid email formats'),
    ('P2 — Next Sprint (2 Weeks)',      'Medium',
     'Inconsistent status values, missing optional fields, outlier amounts'),
    ('P3 — Backlog (1 Month)',          'Low',
     'Casing inconsistencies, minor format deviations'),
]
for pname, sev_label, desc in prios:
    doc_h2(pname)
    doc_bullet(f'Severity: {sev_label}')
    doc_bullet(f'Scope: {desc}')

doc.add_page_break()
ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = ep.add_run('— End of Issues Documentation —\nBusiness Strategy Consulting Division  |  March 2026  |  CONFIDENTIAL')
er.font.name = 'Calibri'; er.font.size = Pt(10); er.font.italic = True; er.font.color.rgb = GRAY

doc.save(f'{OUT}/DQA_Issues_Report.docx')
print(f'  ✅  DQA_Issues_Report.docx saved')
print(f'\n✅  Step 2 complete. Files in: {OUT}')
