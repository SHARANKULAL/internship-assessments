"""
step4_generate_framework.py
Generates the final DQA Framework Word document:
  DQA_Framework_Final.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = r'C:\Users\shara\OneDrive\Desktop\Aptitude\week3\dataset'
os.makedirs(OUT, exist_ok=True)

NAVY  = RGBColor(0x1A, 0x50, 0x9C)
TEAL  = RGBColor(0x00, 0x87, 0x8A)
AMBER = RGBColor(0xF5, 0xA6, 0x23)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GRAY  = RGBColor(0x55, 0x65, 0x7A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def shade_cell(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(18); r.font.bold = True
    r.font.color.rgb = NAVY
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0); rule.paragraph_format.space_after = Pt(8)
    rr = rule.add_run('─' * 85); rr.font.size = Pt(5); rr.font.color.rgb = TEAL

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = TEAL

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = DARK

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = DARK

def bullet(text, sym='▸'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f'{sym}  {text}')
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = DARK

def callout(label, text, bg='E9F1FB', lc=NAVY):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shade_cell(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'{label}  '); r1.font.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = lc
    r2 = p.add_run(text); r2.font.size = Pt(11); r2.font.color.rgb = DARK
    doc.add_paragraph()

def styled_table(headers, rows_data, col_widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    hr = tbl.rows[0]
    for ci, h in enumerate(headers):
        c = hr.cells[ci]; shade_cell(c, '1A509C')
        c.width = Inches(col_widths[ci])
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h); r.font.bold = True; r.font.color.rgb = WHITE
        r.font.name = 'Calibri'; r.font.size = Pt(10)
    for ri, row in enumerate(rows_data):
        dr = tbl.add_row()
        bg = 'E9F1FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            c = dr.cells[ci]; shade_cell(c, bg)
            c.width = Inches(col_widths[ci])
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val)); r.font.name = 'Calibri'; r.font.size = Pt(10)
            r.font.color.rgb = DARK
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
cp = doc.add_paragraph()
cp.paragraph_format.space_before = Pt(60); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cp.add_run('FINALIZED DATA QUALITY FRAMEWORK')
cr.font.name = 'Calibri'; cr.font.size = Pt(26); cr.font.bold = True; cr.font.color.rgb = NAVY
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run('EdTech Platform — Governance, Standards & Continuous Improvement')
sr.font.name = 'Calibri'; sr.font.size = Pt(14); sr.font.italic = True; sr.font.color.rgb = TEAL
rp = doc.add_paragraph(); rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rp.paragraph_format.space_before = Pt(6)
rr2 = rp.add_run('━' * 42); rr2.font.color.rgb = AMBER
mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = mp.add_run('Business Strategy Consulting Division  |  March 2026  |  CONFIDENTIAL')
mr.font.name = 'Calibri'; mr.font.size = Pt(11); mr.font.color.rgb = GRAY
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. FRAMEWORK OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1('1. Framework Overview')
body(
    'This document defines the Finalized Data Quality Framework for EdTech platforms. '
    'It establishes the governance structure, data quality dimensions, validation rule library, '
    'monitoring processes, and continuous improvement mechanisms necessary to sustain high-quality '
    'data across all platform systems and data domains.'
)
body('The framework is built on five interlocking pillars:')
pillars = [
    ('🏛️  Governance & Ownership',  'Clear roles, accountability, and policy structures'),
    ('📐  Standards & Definitions',   'Unified data dictionary, formats, and naming conventions'),
    ('🤖  Validation & Rules',        'Automated checks covering completeness, integrity, format, and logic'),
    ('📊  Monitoring & Alerting',     'Real-time quality scoring, dashboards, and threshold alerts'),
    ('🔁  Continuous Improvement',    'Feedback loops, root-cause resolution, and maturity progression'),
]
for icon_title, desc in pillars:
    bullet(f'{icon_title}: {desc}')

callout('Framework Scope:', 
    'Covers 6 core data entities: Learners, Courses, Assessments, Course Enrollments, '
    'Transactions, and Academic Calendar — the complete EdTech operational data model.',
    bg='E8F6F7', lc=TEAL)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. DATA QUALITY DIMENSIONS
# ═══════════════════════════════════════════════════════════════════════════════
h1('2. Data Quality Dimensions')
body('The framework assesses data across seven quality dimensions. Each dimension maps to specific validation rules applied during automated checks.')
styled_table(
    headers=['Dimension', 'Definition', 'Example Rule', 'Priority'],
    rows_data=[
        ['Completeness', 'Required fields must not be null or empty',
         'learner_id, email, enrollment_date must not be blank', 'Critical'],
        ['Uniqueness',   'Primary keys and unique identifiers must not repeat',
         'No duplicate learner_id, transaction_id, or enrollment_id', 'Critical'],
        ['Validity',     'Values must conform to defined formats and ranges',
         'Email matches regex; grade between 0–100; price ≥ 0', 'High'],
        ['Consistency',  'Values across related records must not conflict',
         'Completion date must be after enrollment date', 'High'],
        ['Referential Integrity', 'Foreign keys must reference existing parent records',
         'Every enrollment must have a valid learner_id and course_id', 'Critical'],
        ['Timeliness',   'Data must be updated within defined SLA windows',
         'Assessment scores updated within 48hrs of submission', 'Medium'],
        ['Accuracy',     'Values must reflect the true real-world state',
         'Transaction amount reconciles with payment gateway receipt', 'High'],
    ],
    col_widths=[1.5, 2.2, 2.8, 0.9]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. GOVERNANCE MODEL
# ═══════════════════════════════════════════════════════════════════════════════
h1('3. Governance Model')
body(
    'Data quality governance distributes accountability across three tiers, '
    'ensuring that no single team owns the entire quality burden.'
)

h2('3.1  Governance Roles & RACI')
styled_table(
    headers=['Role', 'Responsible', 'Accountable', 'Consulted', 'Informed'],
    rows_data=[
        ['Chief Data Officer',      'DQ strategy & policy',     '✅', 'Board, CEO', 'All stakeholders'],
        ['Data Quality Council',    'Monthly review & prioritisation', '✅', 'All stewards', 'Engineering leads'],
        ['Domain Data Stewards',    'Field-level rules, exception handling', '✅', 'BAs, Analysts', 'CDO'],
        ['Data Engineering Team',   'Pipeline validation, monitoring builds', '✅', 'Stewards', 'CDO'],
        ['Business Analysts',       'Dashboard validation, insight sign-off', '✅', 'Stewards', 'Leadership'],
        ['Product Managers',        'Source-system data accuracy', '✅', 'Engineering', 'Stewards'],
    ],
    col_widths=[2.0, 2.5, 0.7, 1.5, 1.7]
)

h2('3.2  Data Quality Council Charter')
bullet('Meeting cadence: Monthly (standing agenda: quality scores, open issues, resolution tracking)')
bullet('Quorum: CDO + at least 3 domain stewards')
bullet('Escalation path: Unresolved Critical issues escalate to CEO within 5 business days')
bullet('Decisions: Quality rule changes, governance exceptions, tool investment')

# ═══════════════════════════════════════════════════════════════════════════════
# 4. VALIDATION RULES LIBRARY
# ═══════════════════════════════════════════════════════════════════════════════
h1('4. Validation Rules Library')
body('The following rules are implemented as automated checks in the data pipeline. Each rule carries a severity level and a defined remediation action.')

tables_rules = {
    'Learners': [
        ('LRN-01', 'learner_id', 'Uniqueness',      'learner_id must be unique',                  'Critical', 'Delete or merge duplicate; reassign FK refs'),
        ('LRN-02', 'email',      'Completeness',     'email must not be null',                     'High',     'Populate from source; mark uncontactable'),
        ('LRN-03', 'email',      'Validity',         'email must match RFC 5322 pattern',           'High',     'Correct or flag for outreach team'),
        ('LRN-04', 'date_of_birth','Validity',       'DOB must not be in the future',              'High',     'Correct to actual date; investigate data entry'),
        ('LRN-05', 'status',     'Consistency',      'status must be Active / Inactive / Pending', 'Low',      'Standardise to Title Case via ETL'),
        ('LRN-06', 'enrollment_date','Validity',     'Enrollment ≥ 2021-01-01 (platform launch)',  'Medium',   'Correct or flag for manual review'),
    ],
    'Courses': [
        ('CRS-01', 'course_title',   'Completeness',  'course_title must not be null',             'High',     'Retrieve from catalogue; suspend if unavailable'),
        ('CRS-02', 'duration_hours', 'Validity',      'duration_hours must be > 0',                'High',     'Set to correct value; minimum 1'),
        ('CRS-03', 'price',          'Validity',      'price must be ≥ 0',                         'High',     'Correct to 0 (free) or actual price'),
        ('CRS-04', 'end_date',       'Consistency',   'end_date must be ≥ start_date',             'Critical', 'Swap dates or correct from schedule'),
        ('CRS-05', 'status',         'Validity',      'status ∈ {Published, Draft, Archived}',     'Low',      'Map typos to valid value via lookup table'),
    ],
    'Assessments': [
        ('ASS-01', 'assessment_id',  'Uniqueness',    'assessment_id must be unique',              'Critical', 'Renumber duplicate; verify with source'),
        ('ASS-02', 'max_score',      'Validity',      'max_score must be > 0',                     'High',     'Set to correct marks value; default 100'),
        ('ASS-03', 'passing_score',  'Business Rule', 'passing_score must be ≤ max_score',         'Critical', 'Set to 50% of max_score as default'),
        ('ASS-04', 'due_date',       'Consistency',   'due_date must be ≥ created_date',           'High',     'Set due_date = created_date + 14 days'),
        ('ASS-05', 'course_id',      'Ref Integrity', 'course_id must exist in courses table',     'Critical', 'Correct or delete orphan assessment'),
    ],
    'Course Enrollments': [
        ('ENR-01', 'enrollment_id',   'Uniqueness',  'enrollment_id must be unique',               'Critical', 'Delete duplicate; keep first occurrence'),
        ('ENR-02', 'learner_id',      'Ref Integrity','learner_id must exist in learners',         'Critical', 'Correct or remove orphan enrollment'),
        ('ENR-03', 'course_id',       'Ref Integrity','course_id must exist in courses',           'Critical', 'Correct or remove orphan enrollment'),
        ('ENR-04', 'completion_date', 'Consistency', 'completion_date ≥ enrollment_date',          'High',     'Correct sequence; add 90-day offset if error'),
        ('ENR-05', 'grade',           'Validity',    'grade must be between 0 and 100',            'High',     'Cap at 100; investigate values above range'),
    ],
    'Transactions': [
        ('TXN-01', 'transaction_id', 'Uniqueness',   'transaction_id must be unique',              'Critical', 'Keep original; mark duplicate as voided'),
        ('TXN-02', 'amount',         'Validity',     'amount must be > 0',                         'Critical', 'If refund, create credit note record'),
        ('TXN-03', 'payment_method', 'Completeness', 'payment_method must not be null',            'Medium',   'Derive from gateway log or mark Unknown'),
        ('TXN-04', 'currency',       'Validity',     'currency must be valid ISO 4217 code',       'Medium',   'Correct to INR / USD / EUR / GBP'),
        ('TXN-05', 'status',         'Validity',     'status ∈ {Completed, Failed, Refunded, Pending}','Medium','Map typos via lookup; default to Completed'),
        ('TXN-06', 'learner_id',     'Ref Integrity','learner_id must exist in learners',          'Critical', 'Investigate; delete if learner cannot be found'),
    ],
    'Academic Calendar': [
        ('CAL-01', 'calendar_id', 'Uniqueness',   'calendar_id must be unique',                    'Critical', 'Renumber duplicate; update FK references'),
        ('CAL-02', 'term_name',   'Completeness', 'term_name must not be null',                    'Medium',   'Assign from institutional schedule'),
        ('CAL-03', 'start_date',  'Completeness', 'start_date must not be null',                   'High',     'Derive from academic_year or institutional doc'),
        ('CAL-04', 'end_date',    'Consistency',  'end_date must be > start_date',                 'Critical', 'Swap or correct from academic schedule'),
        ('CAL-05', 'overlap',     'Business Rule','No two active terms may have overlapping dates', 'High',     'Adjust term boundaries; get sign-off from CDO'),
    ],
}

for tname, rules in tables_rules.items():
    h2(f'4.{list(tables_rules.keys()).index(tname)+1}  {tname}')
    styled_table(
        headers=['Rule ID','Field','Dimension','Rule Description','Severity','Remediation'],
        rows_data=rules,
        col_widths=[0.65, 1.3, 1.0, 2.3, 0.75, 2.3]
    )

# ═══════════════════════════════════════════════════════════════════════════════
# 5. MONITORING & ALERTING
# ═══════════════════════════════════════════════════════════════════════════════
h1('5. Monitoring & Alerting')
body('The monitoring layer ensures data quality is maintained continuously between formal assessment cycles.')

h2('5.1  Quality Scoring Model')
body('Each data entity is assigned a composite quality score (0–100) computed as:')
callout('Formula:', 
    'Quality Score = Σ (Weight × Pass Rate per Dimension)  '
    '| Weighted by severity: Critical=40%, High=30%, Medium=20%, Low=10%',
    bg='E8F6F7', lc=TEAL)

styled_table(
    headers=['Score Band', 'Rating', 'Action Required'],
    rows_data=[
        ['90–100', '🟢 Excellent', 'Maintain; review monthly'],
        ['75–89',  '🟡 Good',      'Investigate top 3 failing rules; resolve within 2 weeks'],
        ['60–74',  '🟠 Fair',      'Priority remediation sprint; CDO review'],
        ['< 60',   '🔴 Poor',      'Immediate escalation; freeze dependent analytics until resolved'],
    ],
    col_widths=[1.2, 1.2, 5.4]
)

h2('5.2  Alert Thresholds')
styled_table(
    headers=['Trigger', 'Threshold', 'Alert Recipient', 'SLA to Resolve'],
    rows_data=[
        ['Critical rule failures',    '≥ 1 occurrence',   'CDO + Engineering Lead', '24 hours'],
        ['Error rate per table',       '> 2%',             'Domain Data Steward',    '48 hours'],
        ['Quality score drop',         '> 10 pts vs prev', 'CDO + Steward',          '72 hours'],
        ['Referential integrity fail', '≥ 1 orphan',       'Engineering Team',       '48 hours'],
        ['Missing required field',     '> 0.5% of records','Domain Steward',         '5 business days'],
    ],
    col_widths=[2.2, 1.5, 2.0, 1.6]
)

h2('5.3  Reporting Cadence')
styled_table(
    headers=['Report', 'Frequency', 'Audience', 'Contents'],
    rows_data=[
        ['Quality Score Dashboard',  'Real-time (daily refresh)', 'All stakeholders', 'Per-table score, trend, top failing rules'],
        ['Steward Weekly Digest',    'Weekly',                    'Data Stewards',    'Open issues, SLA compliance, new failures'],
        ['Council Monthly Report',   'Monthly',                   'DQ Council',       'Score trends, resolution status, roadmap'],
        ['Executive Quality Summary','Quarterly',                 'CEO, Board',       'Quality ROI, maturity progress, investments'],
    ],
    col_widths=[2.0, 1.6, 1.8, 2.9]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. IMPLEMENTATION ROADMAP
# ═══════════════════════════════════════════════════════════════════════════════
h1('6. Implementation Roadmap')
styled_table(
    headers=['Phase', 'Timeline', 'Objective', 'Key Deliverables', 'Success Metric'],
    rows_data=[
        ['1 — Discovery',     'Wks 1–4',   'Baseline current state',
         'Data inventory, quality scorecards, stakeholder map',
         'All 6 tables profiled; quality scores established'],
        ['2 — Framework Design', 'Wks 5–10', 'Define governance & rules',
         'Data dictionary, RACI, validation rules library, tool selection',
         'Governance charter signed; rules library approved by Council'],
        ['3 — Pilot',         'Wks 11–18', 'Deploy on priority tables',
         'Pipeline validation deployed; monitoring dashboard live; steward training',
         'Critical issues < 5; quality score ≥ 70 on Learners & Transactions'],
        ['4 — Scale',         'Wks 19–26', 'Enterprise-wide rollout',
         'All 6 tables monitored; advanced anomaly detection; ROI report',
         'Overall quality score ≥ 85; error rate < 2% platform-wide'],
    ],
    col_widths=[1.2, 0.9, 1.5, 2.5, 2.2]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 7. KPI SCORECARD
# ═══════════════════════════════════════════════════════════════════════════════
h1('7. KPI Scorecard & Success Metrics')
styled_table(
    headers=['KPI', 'Baseline', '6-Month Target', '12-Month Target', 'Owner'],
    rows_data=[
        ['Overall Data Quality Score (0–100)', 'TBD (post-profiling)', '≥ 75', '≥ 90', 'CDO'],
        ['Critical Issue Count per Cycle',     'TBD',                  '< 10', '< 3',  'Engineering'],
        ['Error Rate (all tables)',            'TBD',                   '< 3%', '< 1%', 'Data Stewards'],
        ['Duplicate Record Rate',              'TBD',                   '< 1%', '< 0.3%','Engineering'],
        ['Field Completeness (required)',       'TBD',                  '≥ 97%','≥ 99%', 'Stewards'],
        ['Referential Integrity Pass Rate',     'TBD',                  '≥ 98%','100%',  'Engineering'],
        ['Reporting Accuracy vs Source',        'TBD',                  '≥ 95%','≥ 99%', 'BAs'],
        ['Time to Resolve Critical Issues',     'Ad hoc',               '< 48h','< 24h', 'Engineering'],
        ['Governance Coverage (domains)',       '0%',                   '75%',  '100%',  'CDO'],
        ['Analytics Generation Speed',          'Baseline',             '+30%', '+50%',  'BAs'],
    ],
    col_widths=[2.8, 1.1, 1.3, 1.4, 1.7]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. CONTINUOUS IMPROVEMENT
# ═══════════════════════════════════════════════════════════════════════════════
h1('8. Continuous Improvement Mechanisms')
improvements = [
    ('Root-Cause Analysis Loop',
     'Every Critical issue triggers a root-cause session within 5 days. '
     'Findings feed back into source system design or upstream validation.'),
    ('Rule Enrichment Cycle',
     'Monthly review of new business rules proposed by stewards. '
     'Rules peer-reviewed, tested in sandbox, then promoted to production.'),
    ('Data Quality Maturity Model',
     'Annual maturity assessment against a 5-level scale (Initial → Managed → Defined → Quantified → Optimising). '
     'Progress report presented to board.'),
    ('Vendor & Integration Reviews',
     'Quarterly review of third-party data feeds and integration SLAs. '
     'Enforce contractual data quality standards with LMS and payment providers.'),
    ('Training & Culture Programme',
     'Bi-annual data literacy workshops for all business users. '
     'Champions network across product, finance, and success teams.'),
]
for title, desc in improvements:
    h3(f'▸  {title}')
    body(desc)

callout('Strategic Principle:',
    'Data quality is not a project with an end date — it is a business capability that '
    'must be continuously invested in, measured, and improved to sustain competitive advantage.',
    bg='E8F6F7', lc=TEAL)

# ── Final page ────────────────────────────────────────────────────────────────
doc.add_page_break()
ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = ep.add_run('— End of DQA Framework Document —\nBusiness Strategy Consulting Division  |  March 2026  |  CONFIDENTIAL')
er.font.name = 'Calibri'; er.font.size = Pt(10); er.font.italic = True; er.font.color.rgb = GRAY

doc.save(f'{OUT}/DQA_Framework_Final.docx')
print(f'✅  DQA_Framework_Final.docx saved to: {OUT}')
