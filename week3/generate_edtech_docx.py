
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Colors ────────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x50, 0x9C)
TEAL   = RGBColor(0x00, 0x87, 0x8A)
AMBER  = RGBColor(0xF5, 0xA6, 0x23)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
GRAY   = RGBColor(0x55, 0x65, 0x7A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LTBLUE = RGBColor(0xE9, 0xF1, 0xFB)
LTTEAL = RGBColor(0xE8, 0xF6, 0xF7)

doc = Document()

# ─── Page Margins ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── Style Helpers ──────────────────────────────────────────────────────────────

def set_run_color(run, rgb):
    run.font.color.rgb = rgb

def shade_cell(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Add borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'single'))
        tag.set(qn('w:sz'), kwargs.get('sz', '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '1A509C'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_cover_style():
    """Style document-level defaults."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK

def heading1(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.bold = True
    set_run_color(run, color)
    # bottom border effect via paragraph shading not possible directly;
    # add a thin rule paragraph instead
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after  = Pt(8)
    rule_run = rule.add_run('─' * 80)
    rule_run.font.size = Pt(6)
    set_run_color(rule_run, TEAL)
    return p

def heading2(text, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    set_run_color(run, color)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    set_run_color(run, DARK)
    return p

def bullet(text, level=1, symbol='▸'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.3 * level)
    run = p.add_run(f'{symbol}  {text}')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    set_run_color(run, DARK)
    return p

def callout_box(label, text, bg_hex='E9F1FB', label_color=NAVY):
    """A styled callout / highlight box using a 1-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shade_cell(cell, bg_hex)
    set_cell_border(cell, top='single', left='thick', bottom='single', right='single',
                    sz='8', color='1A509C')
    cell.width = Inches(6.0)
    cp = cell.paragraphs[0]
    cp.paragraph_format.left_indent  = Inches(0.12)
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after  = Pt(4)
    r1 = cp.add_run(f'{label}  ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = label_color
    r2 = cp.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK
    doc.add_paragraph()   # spacing after

def add_page_break():
    doc.add_page_break()

def metric_table(rows_data, headers=None, col_widths=None):
    """Styled table: alternating row colors."""
    n_cols = len(rows_data[0]) if rows_data else 3
    tbl = doc.add_table(rows=0, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    if headers:
        hr = tbl.add_row()
        for ci, hdr in enumerate(headers):
            c = hr.cells[ci]
            shade_cell(c, '1A509C')
            set_cell_border(c, color='1A509C', sz='4')
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(hdr)
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10.5)
    for ri, row in enumerate(rows_data):
        dr = tbl.add_row()
        bg = 'E9F1FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row):
            c = dr.cells[ci]
            shade_cell(c, bg)
            set_cell_border(c, color='CCDDEE', sz='4')
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(cell_text)
            run.font.size = Pt(10.5)
            run.font.color.rgb = DARK
    if col_widths:
        for i, row in enumerate(tbl.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
add_cover_style()

cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(60)
cover.paragraph_format.space_after  = Pt(0)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cover.add_run('DATA QUALITY ASSESSMENT STRATEGY')
cr.font.name = 'Calibri'
cr.font.size = Pt(28)
cr.font.bold = True
cr.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(0)
sr = sub.add_run('for EdTech Platforms')
sr.font.size = Pt(22)
sr.font.color.rgb = TEAL
sr.font.name = 'Calibri'

rule_p = doc.add_paragraph()
rule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rule_p.paragraph_format.space_before = Pt(8)
rule_p.paragraph_format.space_after  = Pt(8)
rr = rule_p.add_run('━' * 40)
rr.font.color.rgb = AMBER

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2r = sub2.add_run('Business Value · Decision-Making Impact · ROI Framework')
s2r.font.size = Pt(13)
s2r.font.italic = True
s2r.font.color.rgb = GRAY
s2r.font.name = 'Calibri'

doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta_p.add_run('Prepared by: Business Strategy Consulting Division\nMarch 2026  |  CONFIDENTIAL')
mr.font.size = Pt(11)
mr.font.color.rgb = GRAY
mr.font.name = 'Calibri'

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
heading1('1. Introduction')
body(
    'The global EdTech industry is undergoing a period of unprecedented expansion. As platforms scale '
    'to serve millions of learners across geographies and learning modalities, data has emerged as the '
    'single most critical strategic asset. Yet for many EdTech organisations, the quality of that data '
    'remains an afterthought — a problem discovered only when dashboards contradict one another, reports '
    'fail audit, or strategic decisions are later revealed to have been built on faulty foundations.'
)
body(
    'Data quality is not merely an IT concern. It is a business imperative. When learner records are '
    'incomplete, assessments are duplicated, or enrolment milestones are recorded inaccurately, the '
    'downstream effects ripple through every layer of the organisation:'
)
bullet('Learning analytics become misleading, causing curriculum teams to invest in the wrong content areas.')
bullet('Student success initiatives target the wrong learner cohorts, wasting resources and reducing retention.')
bullet('Financial reporting incorporates stale or duplicate billing records, creating compliance exposure.')
bullet('Executive dashboards lose the trust of senior leadership, leading to decisions made from instinct rather than evidence.')

body(
    'This document sets out a structured, business-led approach to Data Quality Assessment for EdTech '
    'platforms. It is designed to be understood and championed by business stakeholders — not just '
    'technology teams — because sustainable data quality requires organisational commitment, governance '
    'structures, and clear accountability at every level.'
)

callout_box('Key Insight:', 
    'Organisations with mature data quality programmes report up to 40% faster time-to-insight, '
    '35% reduction in reporting errors, and significantly higher trust in board-level analytics.',
    bg_hex='E8F6F7', label_color=TEAL)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — BUSINESS PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════════════════
heading1('2. Business Problem Statement')
body(
    'Before investing in any framework, it is essential to clearly articulate the problem at hand. '
    'EdTech platforms typically contend with four interconnected data quality challenges that compund '
    'over time as the platform scales:'
)

heading2('2.1  Inconsistent Learner Data')
body(
    'Learner profiles are created and maintained across multiple systems — the Learning Management '
    'System (LMS), the student information system, the CRM, and often multiple third-party course '
    'providers. When these systems are not synchronised, a single learner may appear under different '
    'names, email addresses, or identifiers. Academic history becomes fragmented. Engagement data '
    'points to conflicting patterns depending on which system is queried.'
)
body(
    'The practical impact: personalisation engines recommend content to the wrong cohort; student '
    'success advisors cannot view a complete picture of learner progress; compliance reports contain '
    'conflicting enrolment figures.'
)

heading2('2.2  Duplicate Records')
body(
    'When learners self-register via multiple channels — web portal, mobile app, corporate B2B '
    'partnerships — duplicate accounts proliferate. Without deduplication logic, analytics systems '
    'count the same learner multiple times, inflating active user metrics and distorting conversion '
    'funnel analysis. Leadership makes headcount and growth projections based on figures that do not '
    'reflect reality.'
)

heading2('2.3  Delayed and Stale Data Updates')
body(
    'Assessment scores, completion milestones, and payment statuses are often updated in batch '
    'processes rather than in real time. During the lag window, dashboards display outdated '
    'information. Instructors may believe a learner has not yet completed an assignment when '
    'the completion was recorded hours ago. Finance teams may chase invoices already settled. '
    'These delays erode trust in the platform\'s own data infrastructure.'
)

heading2('2.4  Fragmented Data Systems')
body(
    'Most EdTech platforms have grown through acquisition, rapid feature development, or '
    'integration of best-of-breed SaaS tools. The result is a patchwork of systems — Moodle or '
    'Canvas for course delivery, Salesforce for CRM, Stripe for payments, third-party '
    'assessment engines — each with its own data model, terminology, and update cadence. '
    'There is no unified data dictionary, no single source of truth, and no governance body '
    'to resolve conflicts when the systems disagree.'
)

callout_box('Business Risk Summary:',
    'Data fragmentation is the root cause of analytical inaccuracy. Without resolution, '
    'every decision — from content investment to pricing strategy — carries embedded data risk.',
    bg_hex='E9F1FB', label_color=NAVY)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — MARKET CONTEXT
# ═══════════════════════════════════════════════════════════════════════════════
heading1('3. Market Context')
body(
    'The global EdTech market is projected to surpass $400 billion by 2028, growing at a CAGR '
    'of approximately 14%. This growth is driven by the professionalisation of online learning, '
    'the emergence of enterprise learning platforms, and the increasing expectation of personalised '
    'digital education experiences. However, this growth amplifies the consequences of poor data quality:'
)

bullet('More learners means more data volume — fragmentation risk scales exponentially.')
bullet('Investor and regulatory scrutiny increases with platform size — data accuracy becomes a compliance requirement.')
bullet('Competitive differentiation increasingly hinges on analytics capability — platforms with better data win.')
bullet('Enterprise B2B clients demand reporting-grade data quality as a contractual obligation.')

body(
    'Data-mature EdTech companies — those that have invested in data governance, quality checks, '
    'and integrated data pipelines — demonstrate a measurable competitive advantage. According to '
    'industry research, data-driven EdTech firms achieve:'
)

metric_table(
    [
        ['3× faster',     'Speed of strategic decisions compared to data-reactive peers'],
        ['25% higher',    'Student retention rates from personalisation fuelled by clean data'],
        ['40% reduction', 'In manual reporting effort through validated, automated pipelines'],
        ['60% improvement','In learner outcome prediction accuracy from integrated, clean datasets'],
    ],
    headers=['Metric', 'Business Benefit'],
    col_widths=[1.8, 5.0],
)

body(
    'The implication is clear: data quality is not just an operational hygiene issue. It is a '
    'market positioning strategy. Platforms that embed data quality into their operating model '
    'build a compounding advantage that grows more powerful as data volume increases.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PROPOSED DATA QUALITY STRATEGY
# ═══════════════════════════════════════════════════════════════════════════════
heading1('4. Proposed Data Quality Strategy')
body(
    'The proposed framework is built on four strategic pillars, each addressing a distinct '
    'dimension of data quality. Together they form a governance-led, continuously improving '
    'model that aligns technology capability with business accountability.'
)

heading2('4.1  Data Governance & Ownership')
body(
    'Governance is the organisational foundation. Without clear ownership, quality initiatives '
    'fail to sustain themselves. The governance model assigns explicit accountability across three tiers:'
)
bullet('Chief Data Officer (CDO) or equivalent: Sets policy, arbitrates cross-domain conflicts, reports to CEO/Board.')
bullet('Domain Data Stewards: Own data quality within their system domain (LMS, Finance, Student Success).')
bullet('Data Engineering Team: Implements and maintains automated validation pipelines and monitoring.')
body('A Data Quality Council — a monthly cross-functional forum — reviews quality metrics, prioritises remediation, and governs the data dictionary.')

heading2('4.2  Standardisation Practices')
body(
    'Standardisation eliminates ambiguity. The framework establishes:'
)
bullet('A canonical data dictionary with agreed definitions for all core entities (Learner, Course, Enrolment, Assessment, Transaction).')
bullet('Naming conventions and format standards for all key fields (date formats, identifiers, status codes).')
bullet('A master data management layer that serves as the single source of truth for learner and course records.')

heading2('4.3  Validation Checks')
body(
    'Automated validation rules operate at every ingestion and transformation point:'
)
bullet('Completeness checks: Required fields must not be null (e.g., learner ID, course ID, enrolment date).')
bullet('Uniqueness checks: Deduplication logic flags or merges duplicate learner and course records.')
bullet('Referential integrity checks: Every enrolment record must reference a valid learner and valid course.')
bullet('Range and format checks: Date fields within logical bounds; email fields matching valid format patterns.')
bullet('Timeliness checks: Records updated within defined SLA windows; alerts raised for stale data.')

heading2('4.4  Monitoring Processes')
body(
    'Monitoring transforms quality from a project into an ongoing operational discipline:'
)
bullet('Real-time dashboards surface data quality scores by domain, updated on a defined cadence.')
bullet('Automated alerts trigger on threshold breaches (e.g., error rate exceeds 2% in learner table).')
bullet('Monthly quality reports presented to the Data Quality Council with trend analysis and resolution status.')
bullet('Annual data quality maturity assessment benchmarks progress against industry standards.')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — DATASET STRUCTURE EXAMPLE
# ═══════════════════════════════════════════════════════════════════════════════
heading1('5. Dataset Structure Example')
body(
    'A typical EdTech data environment is built around five core entities. Understanding the '
    'structure of these entities — and the relationships between them — is fundamental to designing '
    'effective quality checks. Below is a representative schema with indicative quality dimensions for each table.'
)

metric_table(
    [
        ['Learners',           'learner_id, name, email, dob, cohort_id, status',
                               'Uniqueness (no duplicates), Completeness (no null IDs/emails), Format (valid email pattern)'],
        ['Courses',            'course_id, title, category, instructor_id, duration_hours, status',
                               'Referential integrity (instructor must exist), Completeness (no null titles)'],
        ['Course Enrolments',  'enrolment_id, learner_id, course_id, enrol_date, completion_date, grade',
                               'Referential integrity (learner and course must exist), Timeliness (completion within duration)'],
        ['Assessments',        'assessment_id, course_id, type, max_score, due_date',
                               'Range checks (max_score > 0), Referential integrity (course must exist)'],
        ['Academic Calendar',  'calendar_id, term_name, start_date, end_date, holidays',
                               'Date ordering (start < end), No overlapping term windows'],
    ],
    headers=['Entity / Table', 'Key Fields', 'Quality Dimensions'],
    col_widths=[1.6, 2.5, 3.0]
)

body(
    'These five entities form the core of learner journey analytics. Data quality issues in any one '
    'of them cascade into downstream reporting errors — which is precisely why governance and '
    'validation must operate at the entity level, not just at the report level.'
)

callout_box('Example Impact:',
    'A missing completion_date in 8% of enrolment records causes completion rate dashboards to '
    'systematically underreport by 8–12%, leading curriculum teams to wrongly deprioritise '
    'high-performing course content.',
    bg_hex='E8F6F7', label_color=TEAL)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — QUALITY ASSESSMENT WORKFLOW
# ═══════════════════════════════════════════════════════════════════════════════
heading1('6. Quality Assessment Workflow')
body(
    'The quality assessment workflow operates as a continuous cycle, not a one-time audit. '
    'Each phase builds on the previous one, progressively improving the reliability of data '
    'across the platform.'
)

heading2('Step 1: Data Profiling')
body(
    'Profiling generates a statistical baseline of the current data landscape. For every core table, '
    'the profiling process captures: completeness rates per field, distinct value counts, '
    'null rates, frequency distributions, and range profiles for numeric and date fields. '
    'The output is a Data Quality Scorecard — a dashboard that assigns a quality score (0–100) '
    'to each entity, enabling business stakeholders to quickly identify the highest-risk data domains.'
)

heading2('Step 2: Validation Rule Execution')
body(
    'Validation rules — defined during framework design — are executed on a scheduled basis '
    'against production data. Each rule produces a pass/fail record logged to the quality '
    'monitoring system. Business analysts review failed-rule summaries to prioritise '
    'remediation efforts. Rules are prioritised by business impact — a failed referential '
    'integrity check on enrolment records is treated as Priority 1; a formatting anomaly in '
    'a rarely-queried metadata field is Priority 3.'
)

heading2('Step 3: Integrity & Consistency Checks')
body(
    'Cross-system consistency checks compare records across different source systems. For example, '
    'the number of unique learner IDs in the LMS should reconcile with the number of active accounts '
    'in the CRM. Discrepancies trigger investigation workflows. Consistency checks are the primary '
    'mechanism for detecting integration failures and data drift between systems.'
)

heading2('Step 4: Monitoring & Continuous Improvement')
body(
    'Monitoring transforms quality from periodic audits into a continuous operational discipline. '
    'Automated monitoring pipelines run on a defined cadence (daily, weekly, or event-triggered). '
    'Quality scores are tracked over time, enabling trend analysis. Deteriorating scores trigger '
    'escalation to the relevant Data Steward. Improving scores are celebrated as evidence of '
    'governance effectiveness — reinforcing the culture of data ownership.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — IMPLEMENTATION ROADMAP
# ═══════════════════════════════════════════════════════════════════════════════
heading1('7. Implementation Roadmap')
body(
    'The implementation follows a phased approach designed to deliver early wins, build '
    'organisational capability progressively, and minimise disruption to ongoing operations.'
)

metric_table(
    [
        ['Phase 1', 'Weeks 1–4',
         'Discovery & Assessment',
         'Data inventory, stakeholder interviews, quality gap analysis, risk scoring baseline, executive sponsor alignment'],
        ['Phase 2', 'Weeks 5–10',
         'Framework Design',
         'Governance charter, data dictionary, quality rules library, KPI definitions, tool selection and architecture design'],
        ['Phase 3', 'Weeks 11–18',
         'Pilot Implementation',
         'Automated checks deployed on priority entities, monitoring dashboards live, data steward training, first quality score report published'],
        ['Phase 4', 'Weeks 19–26',
         'Scale & Optimise',
         'Enterprise-wide rollout, advanced anomaly detection, continuous improvement cycle, ROI measurement and board reporting'],
    ],
    headers=['Phase', 'Timeline', 'Name', 'Key Activities'],
    col_widths=[0.7, 1.0, 1.5, 3.7]
)

body(
    'Key success factors for implementation include: active executive sponsorship, clear '
    'communication about the business rationale for change, dedicated time allocation for '
    'data stewards, and a realistic change management plan recognising that data quality '
    'is a culture shift, not just a technology deployment.'
)

callout_box('Quick Win Strategy:',
    'Focus Phase 1 on the Learners and Course Enrolments tables — the two entities that '
    'most directly impact analytics accuracy. Early wins here build momentum and demonstrate '
    'ROI before the full framework is complete.',
    bg_hex='E9F1FB', label_color=NAVY)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — BUSINESS IMPACT AND ROI
# ═══════════════════════════════════════════════════════════════════════════════
heading1('8. Business Impact and ROI')
body(
    'The business case for data quality investment is compelling and multi-dimensional. '
    'Returns manifest across four primary value dimensions:'
)

heading2('8.1  Improved Analytics Accuracy')
body(
    'Clean, validated data produces analytics that stakeholders trust. When leadership can rely '
    'on dashboards without requiring manual verification, decision velocity increases dramatically. '
    'Curriculum investment decisions, pricing model changes, and student success interventions '
    'are all executed faster and with greater confidence. The estimated reduction in time spent '
    'on data validation and reconciliation across business units is 40–60%.'
)

heading2('8.2  Faster Decision Making')
body(
    'Data quality bottlenecks — waiting for IT to confirm figures, re-running reports with '
    '"corrected" data, resolving discrepancies between system reports — consume significant '
    'management time. Eliminating these bottlenecks returns hundreds of hours per quarter to '
    'productive analytical work. In fast-growing platforms, this speed advantage translates '
    'directly into competitive differentiation.'
)

heading2('8.3  Operational Efficiency')
body(
    'Finance teams close books faster. Student success teams target interventions correctly. '
    'HR systems reconcile headcount without manual cross-checking. Each of these efficiency gains '
    'has a quantifiable cost-saving component. Reduced manual rework, fewer escalations, and '
    'elimination of duplicate-record remediation sprints contribute to an estimated 25% '
    'improvement in operational efficiency across data-dependent functions.'
)

heading2('8.4  Strategic Insights & Competitive Advantage')
body(
    'Platforms with high data quality can credibly offer personalised learning experiences, '
    'demonstrate ROI to enterprise clients with precision, satisfy regulatory reporting '
    'requirements without scrambling, and build investor-grade data rooms for fundraising or '
    'M&A activity. These strategic benefits are difficult to quantify in isolation but represent '
    'the most enduring value of a data quality programme.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — SUCCESS METRICS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('9. Success Metrics')
body(
    'The following KPIs form the measurement framework for the data quality programme. '
    'Baseline values will be established during Phase 1 (Discovery), and targets set for each '
    'subsequent phase milestone.'
)

metric_table(
    [
        ['Data Error Rate',              'Percentage of records failing at least one validation rule',
                                         '< 2% across all core entities', '12-month target'],
        ['Reporting Accuracy Score',     'Dashboard values reconciling with source-of-truth on spot-checks',
                                         '≥ 97% reconciliation rate',      '9-month target'],
        ['Analytics Generation Time',    'Time from data request to dashboard-ready insight',
                                         '50% reduction vs. baseline',     '12-month target'],
        ['Duplicate Record Rate',        'Percentage of learner records flagged as duplicates',
                                         '< 0.5% post-deduplication',      '6-month target'],
        ['Data Completeness Score',      'Average field-level completeness across required fields',
                                         '≥ 98% completeness',             '6-month target'],
        ['Student Insight Accuracy',     'At-risk identification recall vs. actual dropout outcomes',
                                         '30% improvement vs. baseline',   '18-month target'],
        ['Governance Coverage',          'Percentage of data domains with an assigned Data Steward',
                                         '100% coverage',                  '4-month target'],
        ['Compliance Report Pass Rate',  'Regulatory submissions accepted without revision',
                                         '100% first-time pass rate',      '12-month target'],
    ],
    headers=['KPI', 'Definition', 'Target', 'Timeline'],
    col_widths=[1.9, 2.5, 1.8, 1.2]
)

body(
    'These metrics are reviewed monthly by the Data Quality Council and reported quarterly '
    'to the executive team. Progress against targets is the primary measure of programme success.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — RISKS AND MITIGATION
# ═══════════════════════════════════════════════════════════════════════════════
heading1('10. Risks and Mitigation')
body(
    'Every transformation programme carries execution risk. The following risk register '
    'identifies the most significant threats to programme success, with corresponding '
    'mitigation strategies:'
)

metric_table(
    [
        ['Integration Complexity',
         'Legacy systems may not expose data via standard APIs, requiring custom connectors that delay timelines.',
         'High',
         'Conduct a detailed integration feasibility assessment in Phase 1. Prioritise systems with API access and defer legacy connectors to Phase 4.'],
        ['Governance Buy-In',
         'Data Stewards may not prioritise quality tasks alongside operational responsibilities.',
         'High',
         'Embed quality KPIs into Data Steward performance objectives. Allocate dedicated time and back-fill capacity.'],
        ['Data Ownership Disputes',
         'Cross-functional conflicts over who "owns" shared entities (e.g., learner profile across LMS and CRM).',
         'Medium',
         'CDO arbitrates via Data Quality Council. A RACI matrix for each entity is agreed in Phase 2 and ratified by senior leadership.'],
        ['Tool Adoption Resistance',
         'Business analysts may continue using shadow spreadsheets rather than validated data pipelines.',
         'Medium',
         'Change management programme including training, champions network, and executive mandate to use governed data sources.'],
        ['Scope Creep',
         'Stakeholder pressure to include all data domains in Phase 1 risks diluting focus and delaying value.',
         'Medium',
         'Strict scoping with a formal change control process. Phase 1 covers priority entities only; additional scope enters backlog.'],
        ['Data Privacy Compliance',
         'Quality checks may inadvertently expose sensitive learner PII to unauthorised team members.',
         'High',
         'PII masking in all non-production environments. Role-based access control enforced before any profiling begins.'],
    ],
    headers=['Risk', 'Description', 'Likelihood', 'Mitigation'],
    col_widths=[1.5, 2.3, 0.9, 2.7]
)

callout_box('Programme Resilience:',
    'The phased roadmap is the primary risk management mechanism. By delivering value in '
    'incremental phases, the programme builds the organisational trust and momentum required '
    'to navigate these risks without derailing the overall strategy.',
    bg_hex='E8F6F7', label_color=TEAL)

# ─── Final paragraph ─────────────────────────────────────────────────────────
doc.add_paragraph()
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr2 = closing.add_run(
    '— End of Document —\n'
    'Business Strategy Consulting Division  |  March 2026  |  CONFIDENTIAL'
)
cr2.font.size = Pt(10)
cr2.font.italic = True
cr2.font.color.rgb = GRAY
cr2.font.name = 'Calibri'

# ─── Save ────────────────────────────────────────────────────────────────────
output = r'C:\Users\shara\OneDrive\Desktop\Aptitude\week3\EdTech_DataQuality_Report.docx'
doc.save(output)
print(f'✅  Word document saved: {output}')
