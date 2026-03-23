from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── Page Margins ────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ─── Helper: set cell background ─────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top','left','bottom','right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'CCCCCC')
                tcBorders.append(border)
            tcPr.append(tcBorders)

# ─── Helper: heading ─────────────────────────────────────────────────
def add_heading(text, level=1, color="1F3864"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string("2E75B6")
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string("2E75B6")
    return p

# ─── Helper: body paragraph ──────────────────────────────────────────
def add_body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

# ─── Helper: bullet ──────────────────────────────────────────────────
def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

# ─── Helper: add table ───────────────────────────────────────────────
def add_table(headers, rows, header_color="1F3864", alt_row="EBF3FB"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = alt_row if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    set_cell_border(t)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("TELECOM DASHBOARD")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor.from_string("1F3864")

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("END-TO-END PROCESS GUIDE")
r2.bold = True
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor.from_string("2E75B6")

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Built on Qlik Sense | Proof of Concept (POC)")
r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor.from_string("595959")
r3.italic = True

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Prepared for Evaluator Review | March 2026")
r4.font.size = Pt(11)
r4.font.color.rgb = RGBColor.from_string("595959")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 1: PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════
add_heading("1. Project Overview", level=1)
add_body(
    "This document explains the complete end-to-end process of designing, developing, and deploying "
    "an interactive Telecom Analytics Dashboard using Qlik Sense. The dashboard is built for a "
    "fictional telecom company and covers all critical business areas — Revenue, Network, Customer, "
    "Subscription, and Support — using a structured Star Schema data model."
)

add_heading("Objective", level=2)
bullets_overview = [
    "Build a fully interactive Qlik Sense dashboard with 5 analytical sheets",
    "Load and transform raw CSV data using Qlik Sense Load Script (ETL)",
    "Model data using Star Schema with a Link Table to avoid synthetic keys",
    "Apply data quality fixes to handle nulls, type mismatches, and formatting errors",
    "Enable business users to slice data by date, region, plan, and customer segment",
]
for b in bullets_overview:
    add_bullet(b)

add_heading("Technology Stack", level=2)
add_table(
    ["Component", "Tool / Technology"],
    [
        ["Dashboard Platform", "Qlik Sense (Desktop / Cloud)"],
        ["Scripting Language", "Qlik Script (Load Script Editor)"],
        ["Data Format", "CSV Files (Flat Files)"],
        ["Data Model", "Star Schema with Link Table"],
        ["Output", "5 Interactive Dashboard Sheets"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 2: DATA SOURCES
# ═══════════════════════════════════════════════════════════════════════
add_heading("2. Data Sources & Tables", level=1)
add_body(
    "The project uses 8 source tables — 4 Dimension Tables and 4 Fact Tables — each loaded "
    "from separate CSV files. These form the foundation of the Star Schema data model."
)

add_heading("Dimension Tables", level=2)
add_table(
    ["Table Name", "Key Field", "Description"],
    [
        ["DIM_Customer",     "customer_id",     "Customer demographics: name, email, city, credit score, join date, is_active"],
        ["DIM_Plan",         "plan_id",         "Telecom plans: plan name, type (prepaid/postpaid), monthly charge, data limit"],
        ["DIM_Network",      "network_id",      "Network towers: location, technology type (2G/3G/4G/5G), region"],
        ["DIM_Subscription", "subscription_id", "Subscription details: customer-plan mapping, start/end date, status"],
    ]
)

add_heading("Fact Tables", level=2)
add_table(
    ["Table Name", "Key Field", "Description"],
    [
        ["FACT_CDR",     "cdr_id",     "Call Detail Records: call type, duration, data usage, timestamp, network used"],
        ["FACT_Invoice", "invoice_id", "Billing data: invoice date, total amount, paid amount, invoice status"],
        ["FACT_Payment", "payment_id", "Payment transactions: payment method, payment date, amount paid"],
        ["FACT_Support", "ticket_id",  "Support tickets: category, priority, status, open date, resolution time"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 3: ETL PROCESS
# ═══════════════════════════════════════════════════════════════════════
add_heading("3. ETL Process — Qlik Sense Load Script", level=1)
add_body(
    "The Load Script is the heart of the Qlik Sense application. It extracts data from CSV files, "
    "transforms it (cleaning, type conversion, enrichment), and loads it into the in-memory data model. "
    "The script is divided into 7 logical sections."
)

add_heading("Step-by-Step ETL Flow", level=2)

etl_steps = [
    ("Section 1", "Load DIM_Customer",
     "Load customer records. Cast is_active as text ('0'/'1'). Handle nulls in credit_score using NullAsValue."),
    ("Section 2", "Load DIM_Plan",
     "Load plan details. Apply ApplyMap to map plan_type codes to readable labels (PREPAID / POSTPAID)."),
    ("Section 3", "Load DIM_Network",
     "Load network towers with region and technology type. Create a label mapping for 2G/3G/4G/5G."),
    ("Section 4", "Load DIM_Subscription",
     "Load subscriptions. Alias customer_id as sub_CustomerID to avoid key conflicts."),
    ("Section 5", "Load FACT_CDR",
     "Load call records. Fix timestamp format using Timestamp() function. Handle null duration for SMS/MMS."),
    ("Section 6", "Load FACT_Invoice & FACT_Payment",
     "Load billing and payment data. Alias customer_id as pay_CustomerID. Apply payment method ApplyMap."),
    ("Section 7", "Load FACT_Support",
     "Load support tickets. Alias customer_id as sup_CustomerID. Calculate resolution_days from dates."),
    ("Section 8", "Master Calendar",
     "Generate a dynamic date table using AUTOGENERATE. Creates Year, Month, Quarter, Week fields for time intelligence."),
    ("Section 9", "Link Table",
     "Create a centralized LINK_TABLE connecting all fact keys (CustomerID, InvoiceID, etc.) to eliminate synthetic keys."),
]

for sec, title, desc in etl_steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r_sec = p.add_run(f"  {sec}: ")
    r_sec.bold = True
    r_sec.font.size = Pt(10.5)
    r_sec.font.color.rgb = RGBColor.from_string("1F3864")
    r_title = p.add_run(f"{title} — ")
    r_title.bold = True
    r_title.font.size = Pt(10.5)
    r_desc = p.add_run(desc)
    r_desc.font.size = Pt(10.5)

doc.add_paragraph()
add_heading("Key Script Techniques Used", level=2)
techniques = [
    "ApplyMap() — for lookup/enrichment (plan types, payment methods, priority labels)",
    "Interval() / Floor() — to calculate tenure in months from join_date",
    "Timestamp() — to parse and standardize date-time formats",
    "NullAsValue / NullAsNull — to handle missing data explicitly",
    "Alias (AS) — to rename key fields and prevent synthetic keys",
    "AUTOGENERATE + WHILE — to build the Master Calendar dynamically",
]
for t in techniques:
    add_bullet(t)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 4: DATA MODEL
# ═══════════════════════════════════════════════════════════════════════
add_heading("4. Data Modeling — Star Schema with Link Table", level=1)
add_body(
    "The data model follows a Star Schema architecture. Because multiple fact tables share common keys "
    "(customer_id, invoice_id), a centralized Link Table is used to connect all facts without "
    "creating circular references or synthetic keys in Qlik Sense."
)

add_heading("Why Link Table?", level=2)
link_bullets = [
    "Qlik Sense creates 'Synthetic Keys' when two tables share more than one common field",
    "A Link Table collects all shared keys into one place, acting as a hub",
    "All Fact Tables connect to the Link Table via their aliased keys",
    "This ensures clean, correct associations across all data",
]
for b in link_bullets:
    add_bullet(b)

add_heading("Data Model Architecture", level=2)
add_table(
    ["Layer", "Tables", "Connects To"],
    [
        ["Dimension Layer", "DIM_Customer, DIM_Plan, DIM_Network, DIM_Subscription", "LINK_TABLE via primary keys"],
        ["Link Layer",      "LINK_TABLE (CustomerID, InvoiceID, TicketID, etc.)",    "All Fact Tables"],
        ["Fact Layer",      "FACT_CDR, FACT_Invoice, FACT_Payment, FACT_Support",    "LINK_TABLE via aliased keys"],
        ["Time Layer",      "Master Calendar (Date, Year, Month, Quarter, Week)",     "All tables with date fields"],
    ]
)

add_heading("Synthetic Key Resolution", level=2)
add_body("To avoid synthetic keys, customer_id was aliased differently in each fact table:")
add_table(
    ["Table",          "Original Field", "Alias Used"],
    [
        ["DIM_Subscription", "customer_id", "sub_CustomerID"],
        ["FACT_Payment",     "customer_id", "pay_CustomerID"],
        ["FACT_Support",     "customer_id", "sup_CustomerID"],
        ["FACT_Invoice",     "invoice_id",  "inv_InvoiceID (where needed)"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 5: DATA QUALITY FIXES
# ═══════════════════════════════════════════════════════════════════════
add_heading("5. Data Quality Issues — Found & Fixed", level=1)
add_body(
    "Before loading data into the model, 7 data quality issues were identified and resolved "
    "directly in the Load Script to ensure accurate analysis."
)

add_table(
    ["#", "Issue", "Table", "Fix Applied"],
    [
        ["1", "is_active stored as Boolean (True/False) instead of '1'/'0'",   "DIM_Customer",     "Cast using If(is_active=True,'1','0')"],
        ["2", "roaming_flag stored as Boolean instead of text",                  "FACT_CDR",         "Same Boolean-to-string conversion"],
        ["3", "Null duration for SMS/MMS records",                               "FACT_CDR",         "NullAsValue(duration) → set to 0"],
        ["4", "Null data_usage_mb for voice calls",                              "FACT_CDR",         "NullAsValue(data_usage_mb) → set to 0"],
        ["5", "Timestamp format non-standard (e.g., DD/MM/YYYY HH:MM)",          "FACT_CDR",         "Timestamp(call_timestamp, 'DD/MM/YYYY hh:mm')"],
        ["6", "Payment method codes don't match ApplyMap keys",                  "FACT_Payment",     "Standardized codes: 'CC','DC','NB','UPI','CASH'"],
        ["7", "Missing MMS records in call type list",                           "FACT_CDR",         "Added 'MMS' to call type handling logic"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 6: DASHBOARD SHEETS
# ═══════════════════════════════════════════════════════════════════════
add_heading("6. Dashboard Design — 5 Sheets", level=1)
add_body(
    "The dashboard is organized into 5 purposeful sheets, each targeting a specific business domain. "
    "A consistent filter bar (Date, Region, Plan Type) is available across all sheets."
)

# ── Sheet 1
add_heading("Sheet 1: Executive Summary", level=2)
add_body("Purpose: High-level snapshot for management. Shows the overall health of the telecom business.", italic=True)
add_table(
    ["Chart", "Type", "Dimension", "Measure"],
    [
        ["Total Revenue",          "KPI Tile",  "—",             "Sum(total_amount)"],
        ["Active Subscribers",     "KPI Tile",  "—",             "Count({<is_active={'1'}>} CustomerID)"],
        ["Churn Rate %",           "KPI Tile",  "—",             "Count({<is_active={'0'}>}CustomerID)/Count(CustomerID)*100"],
        ["Avg CSAT Score",         "KPI Tile",  "—",             "Avg(csat_score)"],
        ["Revenue by Month",       "Line Chart","Month (Calendar)","Sum(total_amount)"],
        ["Revenue by Plan Type",   "Bar Chart", "plan_type",     "Sum(total_amount)"],
        ["Subscribers by Region",  "Map/Bar",   "region",        "Count(CustomerID)"],
    ]
)

# ── Sheet 2
add_heading("Sheet 2: Revenue & Billing Analysis", level=2)
add_body("Purpose: Deep-dive into invoicing, billing status, and payment trends.", italic=True)
add_table(
    ["Chart", "Type", "Dimension", "Measure"],
    [
        ["Total Invoiced Amount",   "KPI Tile",   "—",               "Sum(total_amount)"],
        ["Total Collected",         "KPI Tile",   "—",               "Sum(paid_amount)"],
        ["Outstanding Balance",     "KPI Tile",   "—",               "Sum(total_amount) - Sum(paid_amount)"],
        ["Invoice Status Breakdown","Pie Chart",  "invoice_status",  "Count(InvoiceID)"],
        ["Revenue Trend",           "Line Chart", "Month",           "Sum(total_amount)"],
        ["Top Paying Customers",    "Bar Chart",  "customer_name",   "Sum(paid_amount)"],
        ["Payment Method Split",    "Donut Chart","payment_method",  "Count(payment_id)"],
    ]
)

# ── Sheet 3
add_heading("Sheet 3: Network & Usage Analytics", level=2)
add_body("Purpose: Analyze call detail records, data consumption, and network performance.", italic=True)
add_table(
    ["Chart", "Type", "Dimension", "Measure"],
    [
        ["Total Calls",             "KPI Tile",  "—",               "Count(cdr_id)"],
        ["Total Data Used (GB)",    "KPI Tile",  "—",               "Sum(data_usage_mb)/1024"],
        ["Avg Call Duration (min)", "KPI Tile",  "—",               "Avg(duration)/60"],
        ["CDR by Call Type",        "Bar Chart", "call_type",       "Count(cdr_id)"],
        ["Data Usage by Tech",      "Bar Chart", "technology_type", "Sum(data_usage_mb)"],
        ["Peak Hour Analysis",      "Heatmap",   "Hour(call_timestamp)", "Count(cdr_id)"],
        ["Network Usage by Region", "Map/Bar",   "region",          "Count(cdr_id)"],
    ]
)

# ── Sheet 4
add_heading("Sheet 4: Customer & Subscription Intelligence", level=2)
add_body("Purpose: Understand customer demographics, churn behavior, and subscription patterns.", italic=True)
add_table(
    ["Chart", "Type", "Dimension", "Measure"],
    [
        ["Total Customers",          "KPI Tile",    "—",                  "Count(CustomerID)"],
        ["Avg Tenure (Months)",      "KPI Tile",    "—",                  "Avg(Floor((Today()-join_date)/30))"],
        ["Churn Rate %",             "KPI Tile",    "—",                  "Count({<is_active={'0'}>}CustomerID)/Count(CustomerID)*100"],
        ["Avg Credit Score",         "KPI Tile",    "—",                  "Avg(credit_score)"],
        ["Churn by Plan Type",       "Bar Chart",   "plan_name",          "Count({<is_active={'0'}>}CustomerID)"],
        ["Subscription Status",      "Pie Chart",   "subscription_status","Count(CustomerID)"],
        ["Tenure Distribution",      "Bar Chart",   "Tenure Bucket (calc)","Count(CustomerID)"],
        ["Credit Score vs Revenue",  "Scatter Plot","CustomerID",         "X=Avg(credit_score), Y=Sum(total_amount)"],
        ["New Subscriptions Trend",  "Line Chart",  "Month (start_date)", "Count(subscription_id)"],
        ["Customer Detail Table",    "Table",       "customer_name, plan_name, city","Sum(total_amount), Avg(credit_score)"],
    ]
)

# ── Sheet 5
add_heading("Sheet 5: Support & Customer Experience", level=2)
add_body("Purpose: Monitor support ticket volumes, resolution efficiency, and SLA compliance.", italic=True)
add_table(
    ["Chart", "Type", "Dimension", "Measure"],
    [
        ["Total Tickets",           "KPI Tile",  "—",              "Count(ticket_id)"],
        ["Avg Resolution Days",     "KPI Tile",  "—",              "Avg(resolution_days)"],
        ["Open Tickets",            "KPI Tile",  "—",              "Count({<ticket_status={'Open'}>}ticket_id)"],
        ["SLA Breach Count",        "KPI Tile",  "—",              "Count({<sla_breached={'Yes'}>}ticket_id)"],
        ["Tickets by Category",     "Bar Chart", "ticket_category","Count(ticket_id)"],
        ["Tickets by Priority",     "Pie Chart", "priority",       "Count(ticket_id)"],
        ["Resolution Time Trend",   "Line Chart","Month",          "Avg(resolution_days)"],
        ["SLA Breach by Category",  "Bar Chart", "ticket_category","Count({<sla_breached={'Yes'}>}ticket_id)"],
        ["Support Agent Table",     "Table",     "agent_name",     "Count(ticket_id), Avg(resolution_days)"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 7: HOW IT ALL CONNECTS
# ═══════════════════════════════════════════════════════════════════════
add_heading("7. End-to-End Flow Summary", level=1)
add_body("Here is how every layer connects from raw data to final dashboard insight:")

flow_steps = [
    ("Step 1", "Raw Data", "CSV files for 8 tables (Customers, Plans, Network, Subscriptions, CDR, Invoices, Payments, Support)"),
    ("Step 2", "Load Script (ETL)", "Data is loaded, cleaned, enriched using ApplyMap, dates standardized, nulls handled"),
    ("Step 3", "Data Model", "Star Schema formed — Dimensions connected to Facts via Link Table. Master Calendar attached."),
    ("Step 4", "Synthetic Key Fix", "customer_id aliased per table (sub_, pay_, sup_) to remove all synthetic keys"),
    ("Step 5", "Dashboard Sheets", "5 sheets built — Executive, Revenue, Network, Customer, Support"),
    ("Step 6", "Interactivity", "Global filters (Date, Region, Plan Type) allow cross-sheet analysis with single click"),
    ("Step 7", "Insight Delivery", "Business users can drill down, filter, and export data from any sheet in real time"),
]

for step, title, desc in flow_steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"🔷 {step} — {title}: ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor.from_string("1F3864")
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)

doc.add_paragraph()
add_heading("Key Business Questions This Dashboard Answers", level=2)
questions = [
    "Which plan type generates the most revenue?",
    "What is the monthly churn trend and which customer segment churns most?",
    "Which network technology has the highest data usage?",
    "What is the outstanding balance across all customers?",
    "Which support ticket categories have the highest SLA breach rate?",
    "Who are the top 10 customers by revenue?",
    "How does credit score correlate with payment behavior?",
]
for q in questions:
    add_bullet(q)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 8: EVALUATOR Q&A PREP
# ═══════════════════════════════════════════════════════════════════════
add_heading("8. Common Evaluator Questions — Prepared Answers", level=1)

qa_pairs = [
    ("Q: Why did you use a Link Table instead of a simple Star Schema?",
     "A: Multiple fact tables (CDR, Invoice, Payment, Support) all share customer_id and invoice_id. "
     "Connecting them directly in Qlik Sense creates Synthetic Keys, which cause incorrect associations. "
     "The Link Table collects all shared keys in one place, acting as a hub — all facts join to it, "
     "eliminating synthetic keys and circular references."),

    ("Q: How did you handle data quality issues in the load script?",
     "A: We identified 7 issues — Boolean/String mismatches in is_active and roaming_flag, null durations "
     "for SMS/MMS records, non-standard timestamp formats, and payment method key mismatches. Each was "
     "fixed directly in the load script using If(), NullAsValue(), Timestamp(), and standardized mapping tables."),

    ("Q: How does the Master Calendar work?",
     "A: The Master Calendar is generated using AUTOGENERATE in a WHILE loop covering the full date range "
     "of the dataset. It creates fields like Year, Month, Quarter, Week, and Day — enabling time-based "
     "analysis across all sheets without needing date fields in every table."),

    ("Q: Why was tenure_months calculated instead of stored?",
     "A: The source data only contains join_date. Tenure is a derived metric calculated as "
     "Floor((Today() - join_date) / 30). This ensures it's always current without needing to update the CSV."),

    ("Q: How is the churn rate calculated?",
     "A: Churn Rate = Count of customers where is_active = '0' divided by total Count of CustomerID, "
     "expressed as a percentage. Set analysis {<is_active={'0'}>} filters churned customers "
     "while the denominator counts all customers."),

    ("Q: How do the 5 sheets relate to each other?",
     "A: All 5 sheets share the same data model and filter bar. A selection made in any filter "
     "(e.g., selecting 'Region = North') instantly propagates across all 5 sheets simultaneously, "
     "enabling cross-domain analysis in a single click."),
]

for q_text, a_text in qa_pairs:
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(8)
    p_q.paragraph_format.space_after = Pt(2)
    rq = p_q.add_run(q_text)
    rq.bold = True
    rq.font.size = Pt(10.5)
    rq.font.color.rgb = RGBColor.from_string("1F3864")

    p_a = doc.add_paragraph()
    p_a.paragraph_format.space_after = Pt(6)
    p_a.paragraph_format.left_indent = Cm(0.5)
    ra = p_a.add_run(a_text)
    ra.font.size = Pt(10.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 9: GLOSSARY
# ═══════════════════════════════════════════════════════════════════════
add_heading("9. Glossary of Terms", level=1)
add_table(
    ["Term", "Definition"],
    [
        ["ETL",             "Extract, Transform, Load — the process of pulling data, cleaning it, and loading it into the model"],
        ["Star Schema",     "A data model with one central fact table surrounded by dimension tables"],
        ["Link Table",      "A Qlik Sense technique to connect multiple fact tables via shared keys without synthetic keys"],
        ["Synthetic Key",   "An automatically generated composite key in Qlik Sense when tables share more than one field"],
        ["ApplyMap()",       "A Qlik Script function that looks up values from a mapping table (like a VLOOKUP)"],
        ["CDR",             "Call Detail Record — a record of every individual call, SMS, or data session"],
        ["Master Calendar", "A generated date table that provides consistent time-based filtering across all sheets"],
        ["CSAT",            "Customer Satisfaction Score — a measure of customer happiness with support resolution"],
        ["SLA",             "Service Level Agreement — a committed response/resolution time for support tickets"],
        ["Churn",           "When a customer stops using the telecom service (is_active = '0')"],
        ["Set Analysis",    "Qlik Sense syntax {<field={'value'}>} used to filter a specific subset within a measure"],
    ]
)

# ─── Save ─────────────────────────────────────────────────────────────
out_path = r"c:\Users\shara\OneDrive\Desktop\Aptitude\POC\Telecom_Dashboard_EndToEnd_Guide.docx"
doc.save(out_path)
print(f"SUCCESS: Document saved to {out_path}")
